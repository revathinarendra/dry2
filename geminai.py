import google.generativeai as genai
from django.conf import settings
from io import BytesIO
import PyPDF2
import docx
import boto3

# Configure Google Generative AI API

genai.configure(api_key=settings.GEMINI_KEY)
gemini_llm = genai.GenerativeModel("gemini-1.5-flash")
gemini_llm = genai.GenerativeModel("gemini-1.5-flash")
model = genai.GenerativeModel("gemini-1.5-flash")

# Function to call Gemini API
def call_gemini(prompt):
    try:
        response = gemini_llm.generate_content(prompt)
        return response.text
    except Exception as e:
        # Log the error and return it
        # You can replace print with a logging call if needed
        print(f"Error calling Gemini LLM: {e}")
        return f"Error calling Gemini LLM: {e}"

# Function to generate job description
def generate_job_description(details):
    prompt = f"""
    Create a professional job description for the role of {details['role']} at {details['job_company_name']}.
    The role requires the following skills: {details['skills']}. The candidate should have the following project experience: {details['project_experience']}.
    Include any other relevant details: {details['other_details']}.
    """
    return call_gemini(prompt)

# Function to generate evaluation criteria
def generate_evaluation_criteria(details, job_description):
    skills_list = details['skills'].split(",")  # Convert comma-separated skills into a list

    # Create the table in Markdown format
    table_header = "| Skill | Evaluation Criteria | Weightage (%) |\n|---|---|---|\n"
    table_rows = ""
    for skill in skills_list:
        table_rows += f"| {skill.strip()} | Evaluate based on proficiency in {skill.strip()} | [Assign appropriate weightage] |\n"

    # Combine job description and table into the final evaluation criteria
    prompt = f"""
    Based on the following job description, create detailed evaluation criteria for assessing candidates for the role of {details['role']} at {details['job_company_name']}.
    
    Job Description:
    {job_description}
    
    Please include the key skills from the job description and assign appropriate weightage (%) to each skill in the following table:
    
    {table_header}{table_rows}
    
    Ensure that the evaluation criteria focuses on the most important aspects of the role and accurately reflects the job requirements outlined in the description.
    """
    return call_gemini(prompt)

def read_resume_type(file):
        """Extract text from PDF or DOCX files."""
        file.seek(0)  # Ensure file pointer is at the beginning
        if file.name.lower().endswith('.pdf'):
            # Wrap file content in BytesIO for safe processing
            pdf_file = BytesIO(file.read())
            reader = PyPDF2.PdfReader(pdf_file)
            resume_text = "".join(page.extract_text() for page in reader.pages)
            pdf_file.close()
            return resume_text
        elif file.name.lower().endswith('.docx'):
            # Wrap file content in BytesIO for safe processing
            docx_file = BytesIO(file.read())
            doc = docx.Document(docx_file)
            resume_text = "\n".join(para.text for para in doc.paragraphs)
            docx_file.close()
            return resume_text
        else:
            return None
def upload_resume_to_cloud(file, file_name, resume_id, content_type):
        """Upload the resume to S3."""
        s3 = boto3.client(
            's3',
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
            region_name=settings.AWS_S3_REGION_NAME
        )

        # Generate a unique file name with the resume_id
        stored_file_name = f"{resume_id}-{file_name}"

        # Upload the file to S3
        s3.upload_fileobj(
            file,
            settings.AWS_STORAGE_BUCKET_NAME,
            stored_file_name,
            ExtraArgs={'ContentType': content_type}
        )

        return stored_file_name


def read_resume(resume_text):
    """Extract fields from resume text."""
    extracted_data = {
        "name": extract_field(resume_text, "name"),
        "mobile": extract_field(resume_text, "mobile"),
        "email": extract_field(resume_text, "email"),
        "role": extract_field(resume_text, "role"),
    }
    return extracted_data


def extract_field(text, field_name):
    """Extract a specific field from the resume text."""
    prompt = f"""The following text is extracted from a resume. Please extract the candidate's {field_name} in one word or short phrase:
    "{text}"
    """
    try:
        extracted_value = call_gemini(prompt=prompt)  # Replace with actual logic
        return extracted_value.strip() if extracted_value else None
    except Exception as e:
        return f"Error: {str(e)}"

def fetch_resume_from_cloud(resume_id):
    s3 = boto3.client(
        's3',
        aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
        aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
        region_name=settings.AWS_S3_REGION_NAME
    )

    # Construct the file name based on the resume_id
    stored_file_name = f"{resume_id}"  # Since resume_id is a UUID, we don't need the timestamp part

    try:
        # Fetch the file from S3
        file_obj = s3.get_object(
            Bucket=settings.AWS_STORAGE_BUCKET_NAME,
            Key=stored_file_name
        )

        # Get the file content and the content type
        file_content = file_obj['Body'].read()
        content_type = file_obj['ContentType']

        # Determine file type based on content type or file extension
        if content_type == 'application/pdf' or stored_file_name.lower().endswith('.pdf'):
            # If PDF, extract text using PyPDF2
            resume_text = read_resume_type(file_content)
        elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or stored_file_name.lower().endswith('.docx'):
            # If DOCX, extract text using python-docx
            resume_text = read_resume_type(file_content)
        elif content_type == 'text/plain' or stored_file_name.lower().endswith('.txt'):
            # If TXT, just decode the content as text
            resume_text = file_content.decode('utf-8')
        else:
            raise Exception(f"Unsupported file type: {content_type}")

        return resume_text

    except Exception as e:
        raise Exception(f"Error fetching resume from cloud: {str(e)}")
def generate_interview_questions(job_description, evaluation_criteria, resume_text):
    prompt = f"""
    Create a set of interview questions based on the following job description, evaluation criteria, and the candidate's resume.

    Consider the following conditions based on the role:

    1. If the role is for freshers or entry-level job seekers:
       - Provide programming questions to evaluate their coding skills, along with a few theoretical questions related to programming concepts.
       - Ask questions about the projects listed in the resume and include any relevant follow-up questions or techniques used in those projects.

    2. If the role is for senior positions or involves architecture:
       - Include relevant programming questions, as well as questions related to system design or solution architecting.

    Job Description:
    {job_description}

    Evaluation Criteria:
    {evaluation_criteria}

    Candidate Resume:
    {resume_text}

    Please provide a tailored list of interview questions that will help assess the candidate's suitability for the role, focusing on the mentioned conditions.
    """
    return call_gemini(prompt)

# Function to read transcription text
def read_transcription(file):
    if file.name.lower().endswith('.txt'):
        return file.read().decode('utf-8')
    elif file.name.lower().endswith('.pdf'):
        reader = PyPDF2.PdfReader(file)
        return "".join(page.extract_text() for page in reader.pages)
    elif file.name.lower().endswith('.docx'):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    return ""



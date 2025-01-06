
import uuid
from flask import Flask, render_template, request, jsonify
import json
import google.generativeai as genai
import time
import os
import chromadb
import chromadb.utils.embedding_functions as embedding_functions
from werkzeug.utils import secure_filename
import sqlite3
from django.conf import settings
import docx
import PyPDF2
from PyPDF2 import PdfReader
import boto3

app = Flask(__name__)
# Configure Google Generative AI API
os.environ["GEMINI_KEY"] = settings.GEMINI_KEY
#genai.configure(settings.GEMINI_KEY)
gemini_llm = genai.GenerativeModel("gemini-1.5-flash")

openai_ef = embedding_functions.OpenAIEmbeddingFunction(
                api_key=settings.API_KEY,
                model_name=settings.MODEL_NAME
            )

# Initialize ChromaDB client
 
headers = {
  'ChromaApiToken': settings.CHROMA_API_TOKEN
}
client = chromadb.HttpClient(host=settings.HOST,headers=headers)



collection = client.get_or_create_collection(name="profiles11",)

# Configure upload folder and allowed extensions
UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# @app.route('/')
# def home():
#     return render_template('home.html')

# Function to interact with Gemini LLM
def call_gemini(prompt):
    try:
        response = gemini_llm.generate_content(prompt)
        return response.text
    except Exception as e:
        # Return the error as a string to the calling function
        return f"Error calling Gemini LLM: {e}"

def generate_job_description(details):
    prompt = f"""
    Create a professional job description for the role of {details['role']} at {details['company_name']}.
    The role requires the following skills: {details['skills']}. The candidate should have the following project experience: {details['project_experience']}.
    Include any other relevant details: {details['other_details']}.
    """
    return call_gemini(prompt)

# @app.route('/generate-job-description', methods=['GET', 'POST'])
def generate_job_description():
    if request.method == 'POST':
        data = request.json

        # Validate required fields
        required_fields = ['company_name', 'role', 'skills', 'project_experience', 'other_details']
        if not all(field in data for field in required_fields):
            return jsonify({"error": "Missing required fields"}), 400

        try:
            # Generate Job Description
            job_description = generate_job_description_helper(data)
            print("Generated Job Description:", job_description)


            # Generate Evaluation Criteria
            evaluation_criteria = generate_evaluation_criteria(data, job_description)

            # # Store data in SQLite database
            # job_id= (
            #     company_name=data['company_name'],
            #     role=data['role'],
            #     job_description=job_description,
            #     evaluation_criteria=evaluation_criteria
            # )
            # job.save()

            return jsonify({
                #"job": job,
                "job_description": job_description,
                "evaluation_criteria": evaluation_criteria
            })
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    #return render_template('generate_job_description.html')

# Helper function to store job data in the database
# def store_in_db(company_name, role, job_description, evaluation_criteria):
#     conn = sqlite3.connect("jobs.db")
#     cursor = conn.cursor()
#     cursor.execute("""
#         INSERT INTO job_descriptions (company_name, role, job_description, evaluation_criteria)
#         VALUES (?, ?, ?, ?)
#     """, (company_name, role, job_description, evaluation_criteria))
#     conn.commit()
#     job_id = cursor.lastrowid  # Get the ID of the newly inserted row
#     conn.close()
#     return job_id


# Helper function to generate evaluation criteria
def generate_evaluation_criteria(details, job_description):
    skills_list = details['skills'].split(",")  # Convert comma-separated skills into a list

    # Create the table in Markdown format
    table_header = "| Skill | Evaluation Criteria | Weightage (%) |\n|---|---|---|\n"
    table_rows = ""
    for skill in skills_list:
        table_rows += f"| {skill.strip()} | Evaluate based on proficiency in {skill.strip()} | [Assign appropriate weightage] |\n"

    # Combine job description and table into the final evaluation criteria
    prompt = f"""
    Based on the following job description, create detailed evaluation criteria for assessing candidates for the role of {details['role']} at {details['company_name']}.

    Job Description:
    {job_description}

    Please include the key skills from the job description and assign appropriate weightage (%) to each skill in the following table:

    {table_header}{table_rows}

    Ensure that the evaluation criteria focuses on the most important aspects of the role and accurately reflects the job requirements outlined in the description.
    """
    return call_gemini(prompt)
# Helper function to generate job description
def generate_job_description_helper(details):
    prompt = f"""
    Create a professional job description for the role of {details['role']} at {details['company_name']}.
    The role requires the following skills: {details['skills']}. The candidate should have the following project experience: {details['project_experience']}.
    Include any other relevant details: {details['other_details']}.
    """
    return call_gemini(prompt)


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def call_gemini_with_file(prompt, file):
    try:
        response = gemini_llm.generate_content([prompt, file])
        return response.text
    except Exception as e:
        return f"Error calling Gemini LLM: {e}"
# for resume field extractions
def generate_uuid():
    """Generate a unique UUID."""
    return str(uuid.uuid4())


def read_resume(resume_file):
    if resume_file.filename.lower().endswith('.pdf'):
        reader = PyPDF2.PdfReader(resume_file)
        resume_text = ""
        for page in reader.pages:
            resume_text += page.extract_text()
        return resume_text
    elif resume_file.filename.lower().endswith('.docx'):
        doc = docx.Document(resume_file)
        resume_text = "\n".join([para.text for para in doc.paragraphs])
        return resume_text
    else:
        return None



# Fetch jobs from SQLite database
def fetch_jobs():
    conn = sqlite3.connect("jobs.db")
    cursor = conn.cursor()
    cursor.execute("SELECT job_id, role, company_name FROM job_descriptions")
    jobs = cursor.fetchall()
    conn.close()
    return jobs
# Routes
@app.route('/find-matching-profiles', methods=['GET', 'POST'])
def find_matching_profiles():
    if request.method == 'POST':
        option = request.form.get("option")

        if option == "add_profiles":
            if 'resumes' not in request.files:
                return jsonify({"error": "No files uploaded"}), 400

            files = request.files.getlist('resumes')
            profiles = []
            count = collection.count()

            for idx, file in enumerate(files):
                if allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    file.save(filepath)

                    # Extract text based on file type
                    resume_text = read_resume(file)
                    if not resume_text:
                        return jsonify({"error": f"Unsupported file format for {filename}"}), 400

                    extracted_fields = {
                        "Name": extract_field(resume_text, "What is the candidate name?"),
                        "Mobile": extract_field(resume_text, "What is the candidate's mobile number?"),
                        "Email": extract_field(resume_text, "What is the candidate's email address?"),
                        "Role": extract_field(resume_text, "What is the candidate's latest role?"),
                        "Text": resume_text
                    }
            
                    #Add to ChromaDB
                    time.sleep(2)
                    collection.add(
                        documents=[extracted_fields["Text"]],
                        metadatas=[{
                            "Name": extracted_fields["Name"],
                            "Mobile": extracted_fields["Mobile"],
                            "Email": extracted_fields["Email"],
                            "Role": extracted_fields["Role"],
                            "Resume_text": extracted_fields["Text"]
                        }],
                        ids=[str(count + idx)]
                    )

                    profiles.append(extracted_fields)

            return jsonify({"message": "Upload successful", "profiles": profiles})

        elif option == "find_profiles":
            job_id = request.form.get("job_id")
            print(job_id)
            if not job_id:
                return jsonify({"error": "No job selected"}), 400

            # Fetch job description and evaluation criteria
            conn = sqlite3.connect("jobs.db")
            cursor = conn.cursor()
            cursor.execute("""
                SELECT job_description, evaluation_criteria FROM job_descriptions WHERE job_id = ?
            """, (job_id,))
            job = cursor.fetchone()
            conn.close()

            if not job:
                return jsonify({"error": "Job not found"}), 404

            job_description, evaluation_criteria = job

            # Perform profile matching
            matched_profiles = search_profiles(job_description)

            return jsonify({
                "message": "Matching profiles found",
                "job_description": job_description,
                "evaluation_criteria": evaluation_criteria,
                "profiles": matched_profiles
            })

    # Fetch available jobs
    conn = sqlite3.connect("jobs.db")
    cursor = conn.cursor()
    cursor.execute("SELECT job_id, role FROM job_descriptions")
    jobs = cursor.fetchall()
    conn.close()
    print(jobs)
    return render_template('find_matching_profiles.html', jobs=jobs)

@app.route('/generate-questions/<name>', methods=['POST'])
def generate_questions(name):
    # Get job description, evaluation criteria, and resume text from the form
    job_description = request.form.get("job_description")
    evaluation_criteria = request.form.get("evaluation_criteria")
    resume_text = request.form.get("resume_text")

    if not all([job_description, evaluation_criteria, resume_text]):
        return jsonify({"error": "Missing required data"}), 400

    # Generate interview questions using a prompt
    questions = generate_interview_questions(job_description, evaluation_criteria, resume_text)

    # Render the template and pass the candidate's name and questions
    return render_template(
        'interview_questions.html',
        candidate_name=name,
        questions=questions
    )

def search_profiles(query, top_n=5):
    collection = client.get_collection(name="profiles11")
    results = collection.query(
        query_texts=[query],
        n_results=top_n
    )
    print(results['metadatas'][0])
    if results['metadatas']:
        matched_profiles = []
        for metadata in results['metadatas'][0]:
            matched_profiles.append({
                "Name": metadata['Name'],
                "Mobile": metadata['Mobile'],
                "Email": metadata['Email'],
                "Role": metadata['Role'],
                "Resume_text": metadata['Resume_text']
            })
        return matched_profiles
    return []

def extract_field(text, field_name):
    # Basic logic for field extraction; replace with your actual extraction logic
    prompt = f"""This text is extracted from a resume , can you answer the questions in one word? . 
    "text" : {text} \n
    "Question": {field_name}
    """
    return call_gemini(prompt=prompt)

def generate_interview_questions(job_description, evaluation_criteria, resume_text):
    prompt = f"""
    Create a set of interview questions based on the following job description, evaluation criteria, and the candidate's resume.

    Consider the following conditions based on the role:

    1. If the role is for freshers or entry-level job seekers:
       - Provide programming questions to evaluate their coding skills, along with a few theoretical questions related to programming concepts.
       - Ask questions about the projects listed in the resume and include any relevant follow-up questions or techniques used in those projects.

    2. If the role is for senior positions or involves architecture:
       - Include relevant programming questions, as well as questions related to system design or solution architecting.

    Job Description:
    {job_description}

    Evaluation Criteria:
    {evaluation_criteria}

    Candidate Resume:
    {resume_text}

    Please provide a tailored list of interview questions that will help assess the candidate's suitability for the role, focusing on the mentioned conditions.
    """
    return call_gemini(prompt)

@app.route('/evaluate-interview', methods=['GET', 'POST'])
def evaluate_interview():
    if request.method == 'POST':
        candidate_name = request.form.get("candidate_name")
        transcription_file = request.files.get("transcription_file")

        if not transcription_file or not allowed_file(transcription_file.filename):
            return "Invalid or missing transcription file", 400

        # Read transcription text
        transcription_text = read_transcription(transcription_file)

        # Prompt for LLM evaluation
        prompt = f"""
        Evaluate the candidate {candidate_name}'s performance in the interview based on the following transcription.
        Assess the candidate on areas such as communication, technical skills, problem-solving, and role-specific knowledge.
        Provide a profile summary and determine if the candidate is suitable for the role or not.

        Transcription:
        {transcription_text}
        """
        evaluation = gemini_llm.generate_content(prompt).text

        # Render the results
        return render_template(
            'evaluation_results.html',
            candidate_name=candidate_name,
            evaluation=evaluation
        )

    # Render the evaluation upload page
    candidate_name = request.args.get("candidate_name", "Unknown Candidate")
    return render_template('evaluate_interview.html', candidate_name=candidate_name)


def read_transcription(file):
    if file.filename.lower().endswith('.txt'):
        return file.read().decode('utf-8')
    elif file.filename.lower().endswith('.pdf'):
        reader = PyPDF2.PdfReader(file)
        return "".join(page.extract_text() for page in reader.pages)
    elif file.filename.lower().endswith('.docx'):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    return ""


if __name__ == "__main__":
    app.run(debug=True)



